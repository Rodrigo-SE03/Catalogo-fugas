import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt
import xlsxwriter
from PIL import Image
import PySimpleGUI as sg
import openpyxl
from openpyxl import load_workbook

sg.theme('Dark')   

layout = [[sg.Text('Selecione o arquivo com as tabelas')],
          [sg.Text('Arquivo', size=(8, 1)), sg.Input(), sg.FileBrowse()],
          [sg.Submit(), sg.Cancel()]]

window = sg.Window('Geração de planilha', layout)
event, values = window.read()

arquivo = values[0]
#num_tabelas = int(values[1])
planilha = r'C:\Users\Rodrigo Santana\Programação\Vazamentos\planilha.xlsx'


documento = Document(arquivo)
num_tabelas = len(documento.tables)

colunas = ['Foto','Local','Elemento/Componente','Classificação','Data','Quantidade','Observações']

print(arquivo)

def ler_tabela(documento,num):
    tabela = documento.tables[num]
    foto = num
    local = tabela.cell(0,2).text
    componente = tabela.cell(1,2).text
    tamanho = tabela.cell(2,2).text
    data = tabela.cell(3,2).text
    quantidade = tabela.cell(4,2).text
    obs = tabela.cell(5,2).text
    
    return [foto,local,componente,tamanho,data,quantidade,obs]


df = pd.DataFrame(columns=colunas)
num = 0
while num<num_tabelas:
    df.loc[(len(df))] = ler_tabela(documento,num)
    num+=1

sheet_name = arquivo.split("/")[-2]

with pd.ExcelWriter(planilha, engine='openpyxl',mode = 'a') as writer:
    df.to_excel(writer, sheet_name=sheet_name)
