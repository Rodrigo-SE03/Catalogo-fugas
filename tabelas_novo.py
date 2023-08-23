import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt
import PySimpleGUI as sg
import re

sg.theme('Dark')   

layout = [[sg.Text('Selecione a planilha')],
          [sg.Text('Planilha', size=(8, 1)), sg.Input(), sg.FileBrowse()],
          [sg.Text('Selecione a pasta de saída')],
          [sg.Text('Pasta', size=(8, 1)), sg.Input(), sg.FolderBrowse()],
          [sg.Text('Selecione a pasta com as fotos')],
          [sg.Text('Pasta', size=(8, 1)), sg.Input(), sg.FolderBrowse()],
          [sg.Submit(), sg.Cancel()]]

window = sg.Window('Geração de tabelas', layout)
event, values = window.read()

planilha = values[0]
pasta_fotos = values[2]
pasta_saida = values[1]
sheet_name = pasta_saida.split("/")[-1]

df = pd.read_excel(planilha,sheet_name = sheet_name)
document = Document()

for index,row in df.iterrows():
    table = document.add_table(rows = 0,cols = 3,style = "Table Grid")
    pic = f"{pasta_fotos}/image{row['Foto']}.png"
    itens = {
             'Local':row["Local"],
             'Elemento/Componente':row["Elemento/Componente"],
             'Classificação':row["Classificação"],
             'Data':row['Data'],
             'Quantidade':row["Quantidade"],
             'Observações':row["Observações"],
             'Nº': row["Foto"]
             }
    if type(itens['Observações']) != str:
        itens['Observações'] = ""
    
    for item in itens.items():
        
        row_cells = table.add_row().cells
        p1 = row_cells[1].paragraphs[0]
        paragraph_format = p1.paragraph_format
        paragraph_format.line_spacing = 2.0
        p1.add_run(str(item[0])).bold = True

        paragraph_format.space_after
        p2 = row_cells[2].paragraphs[0]
        paragraph_format = p2.paragraph_format
        paragraph_format.line_spacing = 1.0
        p2.add_run(str(item[1]))
        #row_cells[2].text = str(item[1])

    i = 1
    photo_col = table.column_cells(0)
    photo_cell = photo_col[0]
    while i<7:
        photo_cell.merge(photo_col[i])
        i+=1
    photo_cell.text = ""
    paragraph = photo_cell.paragraphs[0]
    run = paragraph.add_run()
    try:
        run.add_picture(pic,1500000)
    except:
        try:
            pic = f"{pasta_fotos}/image{row['Foto']}.jpeg"
            run.add_picture(pic,1500000)
        except:
            try:
                pic = f"{pasta_fotos}/image{row['Foto']}.jpg"
                run.add_picture(pic,1500000)
            except:
                pass

    i=0
    while i<7:
        j=0
        while j<3:
            table.cell(i,j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            j+=1
        i+=1
    
    p = document.add_paragraph()
    
    

document.save(f'{pasta_saida}/tabelas_ordenadas.docx')


#depois que o programa estiver melhor, coloque a interface gráfica e permita inserir os valores das tarifas, horas de trabalho e dias de trabalho por mês