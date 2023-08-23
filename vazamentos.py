import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt
import PySimpleGUI as sg
import re

sg.theme('Dark')   

layout = [[sg.Text('Selecione a pasta que contem pasta de fotos e a planilha')],
          [sg.Text('Pasta', size=(8, 1)), sg.Input(), sg.FolderBrowse()],
          [sg.Submit(), sg.Cancel()]]

window = sg.Window('Geração de tabelas', layout)
event, values = window.read()

pasta = values[0]
df = pd.read_excel(open(f'{pasta}/vazamentos_offline.xlsx','rb'))
document = Document()
qtd_total = 0
qtd_pequeno = 0
qtd_medio = 0
qtd_grande = 0
qtd_extragrande = 0

for index,row in df.iterrows():
    table = document.add_table(rows = 0,cols = 3,style = "Table Grid")
    pic = f"{pasta}/Fotos - Vazamentos/{row['Foto']}.png"
    itens = {
             'Local':row["Local"],
             'Elemento/Componente':row["Componente"],
             'Classificação':row["Classificação"],
             'Data':row['Data'],
             'Quantidade':row["Quantidade"],
             'Observações':row["Observações"]
             }
    if type(itens['Observações']) != str:
        itens['Observações'] = ""
    
    data = re.search(r'([0-9]+)[\-]([0-9]+)[\-]([0-9]+)',str(itens['Data'])).groups()

    qtd_total = qtd_total + itens['Quantidade']
    if itens['Classificação'] == "Pequeno":
        qtd_pequeno = qtd_pequeno + itens['Quantidade']
    elif itens['Classificação'] == "Médio":
        qtd_medio = qtd_medio + itens['Quantidade']
    elif itens['Classificação'] == "Grande":
        qtd_grande = qtd_grande + itens['Quantidade']
    else:
        qtd_extragrande = qtd_extragrande + itens['Quantidade']

    itens['Data'] = f"{data[2]}/{data[1]}/{data[0]}"
    
    for item in itens.items():
        
        row_cells = table.add_row().cells
        p1 = row_cells[1].paragraphs[0]
        paragraph_format = p1.paragraph_format
        paragraph_format.line_spacing = 2.0
        p1.add_run(str(item[0])).bold = True

        paragraph_format.space_after
        p2 = row_cells[2].paragraphs[0]
        paragraph_format = p2.paragraph_format
        paragraph_format.line_spacing = 1.0
        p2.add_run(str(item[1]))
        #row_cells[2].text = str(item[1])

    i = 1
    photo_col = table.column_cells(0)
    photo_cell = photo_col[0]
    while i<6:
        photo_cell.merge(photo_col[i])
        i+=1
    photo_cell.text = ""
    paragraph = photo_cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(pic,1500000)

    i=0
    while i<6:
        j=0
        while j<3:
            table.cell(i,j).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            j+=1
        i+=1
    
    p = document.add_paragraph()
    
    '''
    cont = 0
    if cont == 2:
        p.add_run("\n\n\n\n\n")
        cont = -1
    cont+=1
    '''
    

def analise_custos(qtd_total,tarifa_p_ml,tarifa_p_c,tarifa_fp_ml,tarifa_fp_c,hp,hfp,dias_p_mes):
    tp = tarifa_p_c+tarifa_p_ml
    tfp = tarifa_fp_ml+tarifa_fp_c
    desp_energia = qtd_total*0.3*(hp+hfp)*dias_p_mes*12
    custo = qtd_total*0.3*(hp*tp+hfp*tfp)*dias_p_mes*12

    return[desp_energia,custo]

#print(analise_custos(qtd_total,1.56898,0.85865,0.148332,0.14833,2,14,22))

print(f"Total: {qtd_total}\nPequenos: {qtd_pequeno}\nMedios:{qtd_medio}\nGrandes:{qtd_grande}\nExtragrandes: {qtd_extragrande}")
document.save(f'{pasta}/tabelas.docx')


#depois que o programa estiver melhor, coloque a interface gráfica e permita inserir os valores das tarifas, horas de trabalho e dias de trabalho por mês